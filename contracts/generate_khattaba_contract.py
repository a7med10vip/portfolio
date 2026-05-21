"""
Generates the Khattaba project development contract as a .docx file.
Output: khattaba-contract-v1.docx in the same directory.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy
import os

# ============== BRAND COLORS ==============
GREEN = RGBColor(0x30, 0xC2, 0x80)
DARK = RGBColor(0x0A, 0x0A, 0x0A)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)
LIGHT_GRAY = RGBColor(0xE5, 0xE5, 0xE5)
PLACEHOLDER = RGBColor(0xC0, 0x39, 0x2B)  # red for [...] placeholders
GREEN_BG = "F0FBF6"
DARK_BG = "0A0A0A"

FONT_AR = "Sakkal Majalla"
FONT_AR_ALT = "Tahoma"

# ============== HELPERS ==============

def set_rtl(paragraph):
    """Force right-to-left direction on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def set_cell_rtl(cell):
    for p in cell.paragraphs:
        set_rtl(p)


def set_arabic_font(run, size=12, bold=False, color=None, font=FONT_AR):
    """Apply Arabic-compatible font + size to a run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Set both ASCII and complex script font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)  # complex script (Arabic)
    rFonts.set(qn("w:eastAsia"), font)
    # complex script size
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    # mark as RTL run for bidi rendering
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    if bold:
        bcs = OxmlElement("w:bCs")
        rPr.append(bcs)


def add_arabic_para(doc, text, size=12, bold=False, color=None,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4,
                    space_before=0, line_spacing=1.5, indent_first=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if indent_first is not None:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    run = p.add_run(text)
    set_arabic_font(run, size=size, bold=bold, color=color)
    return p


def add_arabic_para_runs(doc, segments, size=12, bold=False,
                         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4,
                         space_before=0, line_spacing=1.5):
    """Mixed-format paragraph from list of (text, format_overrides)."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    for seg in segments:
        text = seg["text"]
        s = seg.get("size", size)
        b = seg.get("bold", bold)
        c = seg.get("color", None)
        run = p.add_run(text)
        set_arabic_font(run, size=s, bold=b, color=c)
    return p


def add_heading_article(doc, number, title):
    """Article heading: 'المادة X: title' in green bold large."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"المادة {number}: ")
    set_arabic_font(run, size=16, bold=True, color=GREEN)
    run2 = p.add_run(title)
    set_arabic_font(run2, size=16, bold=True, color=DARK)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "30C280")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_arabic_font(run, size=13, bold=True, color=DARK)
    return p


def add_numbered_item(doc, num, text, indent=0.5):
    """Arabic numbered list item: 1- text."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(indent)
    p.paragraph_format.left_indent = Cm(indent * 2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    num_run = p.add_run(f"{num}- ")
    set_arabic_font(num_run, size=12, bold=True, color=GREEN)
    body_run = p.add_run(text)
    set_arabic_font(body_run, size=12, color=DARK)
    return p


def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(indent)
    p.paragraph_format.left_indent = Cm(indent * 2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    bullet_run = p.add_run("•  ")
    set_arabic_font(bullet_run, size=12, bold=True, color=GREEN)
    body_run = p.add_run(text)
    set_arabic_font(body_run, size=12, color=DARK)
    return p


def add_placeholder_para(doc, label, placeholder):
    """Label: ____placeholder____ — placeholder in red."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    lbl = p.add_run(f"{label}: ")
    set_arabic_font(lbl, size=12, bold=True, color=DARK)
    ph = p.add_run(f"[ {placeholder} ]")
    set_arabic_font(ph, size=12, bold=True, color=PLACEHOLDER)
    return p


def add_kv_line(doc, label, value, indent=0.0):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    lbl = p.add_run(f"{label}: ")
    set_arabic_font(lbl, size=12, bold=True, color=DARK)
    val = p.add_run(value)
    set_arabic_font(val, size=12, color=DARK)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def make_table_rtl(table):
    """Reverse cell visual order so it reads right-to-left."""
    tblPr = table._tbl.tblPr
    bidiVisual = OxmlElement("w:bidiVisual")
    tblPr.append(bidiVisual)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def write_cell(cell, text, *, size=11, bold=False, color=None,
               align=WD_ALIGN_PARAGRAPH.RIGHT, shade=None):
    if shade:
        shade_cell(cell, shade)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_arabic_font(run, size=size, bold=bold, color=color)


def add_section_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break()


def add_brand_separator(doc):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("◆  ◆  ◆")
    set_arabic_font(run, size=10, color=GREEN)


def add_callout(doc, label, body, color_hex="30C280"):
    """Single-row table acting as a callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_table_rtl(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FBF9")
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "E0E8E4")
        tcBorders.append(b)
    right = OxmlElement("w:right")
    right.set(qn("w:val"), "single")
    right.set(qn("w:sz"), "24")
    right.set(qn("w:color"), color_hex)
    tcBorders.append(right)
    tcPr.append(tcBorders)

    p1 = cell.paragraphs[0]
    set_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label)
    set_arabic_font(r1, size=10, bold=True, color=RGBColor.from_string(color_hex))

    p2 = cell.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_arabic_font(r2, size=11, color=DARK)

    # spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ============== BUILD DOCUMENT ==============

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        # RTL page direction at section level
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)
        # Page numbers footer (basic)
    # Default Normal style font
    normal = doc.styles["Normal"]
    normal.font.name = FONT_AR
    normal.font.size = Pt(12)

    # ============== COVER ==============
    cover_spacer = doc.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(60)

    add_arabic_para(doc, "بسم الله الرحمن الرحيم",
                    size=14, bold=True, color=GREEN,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=40)

    add_arabic_para(doc, "عقد تطوير وتنفيذ",
                    size=20, bold=True, color=DARK,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=8)

    add_arabic_para(doc, "منصة خطابة السعودية الأولى",
                    size=28, bold=True, color=GREEN,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=8)

    add_arabic_para(doc, "موقع إلكتروني · تطبيق جوال · لوحة تحكم إدارية",
                    size=12, color=GRAY,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=60)

    # Brand separator (large)
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(40)
    run = sep.add_run("◆")
    set_arabic_font(run, size=20, color=GREEN)

    # Parties summary table on cover
    add_arabic_para(doc, "أطراف العقد",
                    size=13, bold=True, color=DARK,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=8)

    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    make_table_rtl(t)
    set_table_borders(t)
    t.columns[0].width = Cm(8)
    t.columns[1].width = Cm(8)
    write_cell(t.cell(0, 0), "الطرف الأول · العميل",
               size=10, bold=True, color=GREEN,
               align=WD_ALIGN_PARAGRAPH.CENTER, shade=GREEN_BG)
    write_cell(t.cell(0, 1), "الطرف الثاني · المطوّر",
               size=10, bold=True, color=GREEN,
               align=WD_ALIGN_PARAGRAPH.CENTER, shade=GREEN_BG)
    write_cell(t.cell(1, 0),
               "شركة خطابة السعودية الأولى\nللتجارة (ذ.م.م)",
               size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(t.cell(1, 1),
               "أحمد علي بسيوني علي\n(مطوّر مستقل)",
               size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Doc info at bottom
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    add_arabic_para_runs(doc, [
        {"text": "رقم العقد: ", "bold": True, "color": DARK},
        {"text": "[ يُحدد لاحقاً ]", "bold": True, "color": PLACEHOLDER},
    ], size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_arabic_para_runs(doc, [
        {"text": "تاريخ التحرير: ", "bold": True, "color": DARK},
        {"text": "[ يُحدد عند التوقيع ]", "bold": True, "color": PLACEHOLDER},
    ], size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_arabic_para_runs(doc, [
        {"text": "محل التحرير: ", "bold": True, "color": DARK},
        {"text": "المملكة العربية السعودية", "color": DARK},
    ], size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Page break to first article
    doc.add_page_break()

    # ============== PARTIES (الديباجة) ==============
    add_arabic_para(doc, "الديباجة وأطراف العقد",
                    size=18, bold=True, color=GREEN,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=18)

    # Party 1
    add_subheading(doc, "الطرف الأول · العميل")
    add_kv_line(doc, "الاسم", "شركة خطابة السعودية الأولى للتجارة (ذات مسؤولية محدودة)")
    add_kv_line(doc, "السجل التجاري", "1010158509")
    add_placeholder_para(doc, "تاريخ القيد", "يُملأ من قِبل الشركة")
    add_placeholder_para(doc, "الرقم الضريبي (VAT)", "إن وُجد")
    add_placeholder_para(doc, "المقر الرئيسي", "العنوان الوطني الكامل · المدينة · الرمز البريدي")
    add_placeholder_para(doc, "الممثل القانوني المفوّض بالتوقيع", "الاسم الرباعي")
    add_placeholder_para(doc, "صفة الممثل", "مدير عام · شريك · وكيل بصلاحية…")
    add_placeholder_para(doc, "رقم هوية الممثل", "—")
    add_placeholder_para(doc, "البريد الإلكتروني الرسمي", "—")
    add_placeholder_para(doc, "رقم الجوال والواتساب الرسمي", "—")
    add_arabic_para(doc, "ويُشار إليه فيما بعد بـ «الطرف الأول» أو «العميل».",
                    size=11, color=GRAY, space_before=4, space_after=12,
                    line_spacing=1.4)

    # Party 2
    add_subheading(doc, "الطرف الثاني · المطوّر")
    add_kv_line(doc, "الاسم الرباعي", "أحمد علي بسيوني علي")
    add_kv_line(doc, "الصفة", "مطوّر برمجيات مستقل")
    add_kv_line(doc, "الجنسية", "مصري")
    add_kv_line(doc, "رقم الهوية الوطنية المصرية", "30309211501095")
    add_kv_line(doc, "محل الإقامة",
                "جمهورية مصر العربية · محافظة كفر الشيخ · شارع جيهان بجوار قسم شرطة ثانٍ")
    add_kv_line(doc, "البريد الإلكتروني الرسمي", "hello@ahmedali.online")
    add_kv_line(doc, "رقم الجوال والواتساب", "+201011648156")
    add_arabic_para(doc, "ويُشار إليه فيما بعد بـ «الطرف الثاني» أو «المطوّر».",
                    size=11, color=GRAY, space_before=4, space_after=12,
                    line_spacing=1.4)

    # ============== التمهيد ==============
    add_subheading(doc, "تمهيد")
    add_arabic_para(doc,
        "حيث إن الطرف الأول جهةٌ تجاريةٌ مرخّصة، تعمل في مجال الوساطة في الزواج "
        "الشرعي داخل المملكة العربية السعودية، ويرغب في تطوير منصة إلكترونية متكاملة "
        "(موقع ويب + تطبيق جوال + لوحة تحكم إدارية) تخدم أعضاءه الجدد والقدامى وتنظّم "
        "عمليات التواصل والمتابعة وفقاً لرؤيته التشغيلية المعتمدة.",
        size=12, space_after=6, line_spacing=1.6, indent_first=0.5)

    add_arabic_para(doc,
        "وحيث إن الطرف الثاني مطوّر برمجيات يمتلك الخبرة الفنية اللازمة لتنفيذ "
        "المشروع الموصوف في الوثيقة التقنية المرفقة بهذا العقد (الملحق أ)، وقد "
        "أبدى استعداده وقبوله للتنفيذ وفق المواصفات والمدد والقيمة المتفق عليها.",
        size=12, space_after=6, line_spacing=1.6, indent_first=0.5)

    add_arabic_para(doc,
        "وحيث اطّلع الطرفان على الوثيقة التقنية المرفقة كاملةً، وفهما مضمونها فهماً "
        "نافياً للجهالة، وأقرّا بأنها تمثّل نطاق العمل الكامل للمشروع.",
        size=12, space_after=6, line_spacing=1.6, indent_first=0.5)

    add_arabic_para(doc,
        "وبناءً على ما تقدّم، وبعد أن أقرّ كلا الطرفين بكامل أهليتهما المعتبرة "
        "شرعاً ونظاماً، اتفقا على ما يلي:",
        size=12, bold=True, space_after=6, line_spacing=1.6, indent_first=0.5)

    doc.add_page_break()

    # ============== المادة 1: التعريفات ==============
    add_heading_article(doc, "الأولى", "التعريفات")
    add_arabic_para(doc,
        "تكون للكلمات والعبارات التالية، أينما وردت في هذا العقد، المعاني الموضّحة قرين كلٍّ منها:",
        size=12, space_after=8, line_spacing=1.5)

    defs = [
        ("المشروع", "منصة خطابة السعودية الأولى بجميع مكوّناتها التقنية (الموقع الإلكتروني، تطبيق الجوال لنظامي iOS و Android، لوحة التحكم الإدارية، خدمات الواجهة الخلفية، قاعدة البيانات)."),
        ("الوثيقة التقنية / الملحق (أ)", "وثيقة نطاق العمل الكاملة المرفقة بهذا العقد، والتي تتضمن اثنا عشر قسماً تفصيلياً واثنين وعشرين متطلباً وميزة، وتُعدّ جزءاً لا يتجزأ منه."),
        ("المنصة", "الموقع الإلكتروني المُطوَّر للعميل، المتاح عبر اسم النطاق المعتمد، بما يشمل جميع صفحاته العامة والخاصة."),
        ("التطبيق", "تطبيق الجوال المُطوَّر بتقنية Flutter لنظامي iOS و Android، المتصل بنفس واجهات البرمجة وقاعدة البيانات الخاصة بالمنصة."),
        ("لوحة التحكم", "الواجهة الإدارية المخصّصة لإدارة المنصة، بما تشمله من شاشات مراجعة الطلبات وإدارة الأعضاء والمحادثات والتقارير."),
        ("المخرجات", "كل ما يُسلَّم للعميل من تصاميم وكود مصدري ووثائق تدريبية وتقارير ضمن نطاق هذا العقد."),
        ("المراحل", "المراحل الخمس المتتالية المنصوص عليها في المادة الثالثة وجدول الدفعات في المادة الخامسة."),
        ("الاعتماد", "الموافقة الكتابية الصريحة من الطرف الأول على مخرجات مرحلة معيّنة، إما عبر البريد الإلكتروني الرسمي أو عبر الواتساب الرسمي للممثل المفوّض."),
        ("الكود المصدري", "ملفات البرمجة الكاملة للمشروع المخزّنة في مستودع GitHub مخصّص للعميل، مع جميع التبعيات وملفات الإعداد اللازمة لإعادة بناء النظام."),
        ("بيئة الإنتاج", "البيئة التشغيلية الفعلية المنشورة على الاستضافة السعودية المعتمدة من الطرف الأول، والتي يصلها المستخدمون النهائيون."),
        ("فترة الضمان", "ثلاثة أشهر تبدأ من تاريخ الإطلاق الرسمي للمنصة والتطبيق، يلتزم خلالها الطرف الثاني بإصلاح الأخطاء التقنية في نطاق العمل المعتمد دون مقابل."),
        ("الجهة الخارجية", "أي مزوّد خدمة طرف ثالث يُستخدم في المشروع (مثل: بوابة الدفع، مزوّد الـ OTP، خدمات الاستضافة، خدمات الإشعارات)."),
        ("نطاق العمل الإضافي", "أي عمل يقع خارج ما هو موصوف في الملحق (أ)، ويُعامَل وفقاً لأحكام المادة الحادية عشرة."),
    ]

    for term, definition in defs:
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.45
        r1 = p.add_run(f"«{term}»: ")
        set_arabic_font(r1, size=12, bold=True, color=GREEN)
        r2 = p.add_run(definition)
        set_arabic_font(r2, size=12, color=DARK)

    # ============== المادة 2: موضوع العقد ==============
    add_heading_article(doc, "الثانية", "موضوع العقد")
    add_arabic_para(doc,
        "يلتزم الطرف الثاني بتطوير وتسليم المشروع للطرف الأول وفق المواصفات "
        "الفنية والنطاق الكامل الموصوف في الوثيقة التقنية المرفقة (الملحق أ)، "
        "والذي يشمل بشكل جوهري ما يلي:",
        size=12, space_after=6, line_spacing=1.5, indent_first=0.5)

    add_numbered_item(doc, 1, "موقع إلكتروني (Next.js) متجاوب يضمّ سبع عشرة صفحة عامة وخاصة بدعم كامل للغة العربية واتجاه RTL.")
    add_numbered_item(doc, 2, "تطبيق جوال (Flutter) لنظامَي iOS و Android، يضمّ كافة الشاشات بتصميم محسّن للأجهزة المحمولة، ومرتبط بنفس قاعدة البيانات وواجهات البرمجة.")
    add_numbered_item(doc, 3, "لوحة تحكم إدارية متكاملة تتضمن تسع شاشات لإدارة الطلبات والأعضاء والمحادثات والتقارير والإعدادات والصلاحيات.")
    add_numbered_item(doc, 4, "الخدمات الخلفية وقاعدة البيانات وواجهات البرمجة، وربط نظام الدفع وخدمات الإشعارات والـ OTP.")
    add_numbered_item(doc, 5, "ترحيل بيانات الأعضاء النشطين من النظام القديم إلى النظام الجديد.")
    add_numbered_item(doc, 6, "اختبارات الجودة والأمان والأداء، ونشر المنصة والتطبيق في بيئة الإنتاج.")

    add_arabic_para(doc,
        "وتُعدّ الوثيقة التقنية المرفقة جزءاً لا يتجزأ من هذا العقد، وتقرأ معه قراءة واحدة.",
        size=11, color=GRAY, space_before=6, space_after=4, line_spacing=1.5)

    # ============== المادة 3: المراحل والمخرجات ==============
    add_heading_article(doc, "الثالثة", "المراحل والمخرجات")
    add_arabic_para(doc,
        "يُنفَّذ المشروع على خمس مراحل متتالية، ولكل مرحلة مخرجات محدّدة "
        "تخضع للاعتماد الكتابي من الطرف الأول قبل الانتقال للمرحلة التالية:",
        size=12, space_after=8, line_spacing=1.5, indent_first=0.5)

    phases = [
        ("الأولى · التصميم (UI/UX)", "أسبوعان · الأسابيع 1-2",
         "تصميم كامل لصفحات الموقع ولوحة التحكم وشاشات التطبيق بصيغة Figma، تصميم متجاوب يعمل على كل أحجام الشاشات، اعتماد الهوية البصرية، اعتماد العميل النهائي على التصميم."),
        ("الثانية · الواجهة الأمامية ولوحة التحكم", "ثلاثة أسابيع · الأسابيع 3-5",
         "تحويل التصميم إلى كود Next.js كامل الاستجابة، بناء جميع صفحات الموقع ولوحة التحكم بتقنية Server-Side Rendering، تهيئة دعم RTL، تهيئة الـ SEO الأساسية."),
        ("الثالثة · الخدمات الخلفية وربط الأنظمة", "ثلاثة أسابيع · الأسابيع 5-7",
         "بناء قاعدة البيانات وواجهات البرمجة، ربط نظام المصادقة و OTP وبوابة الدفع، بناء نظام الشات المراقب ونظام الإشعارات، ترحيل بيانات الأعضاء، إعداد سجل المراجعة، تطبيق متطلبات PDPL التقنية."),
        ("الرابعة · تطبيق الجوال (Flutter)", "ثلاثة أسابيع · الأسابيع 7-10",
         "تطوير التطبيق الكامل لنظامي iOS و Android، ربط الإشعارات عبر Firebase Cloud Messaging، إتاحة الشات المراقب داخل التطبيق، دعم كامل لاتجاه RTL وتحسين الأداء للأجهزة منخفضة المواصفات."),
        ("الخامسة · الاختبار والإطلاق", "أسبوعان · الأسبوعان 11-12",
         "اختبار شامل وظيفي وأمني، اختبار التوافق مع المتصفحات والأجهزة، اختبار الأداء تحت الضغط، إعداد بيئة الإنتاج، إعادة التوجيه من النطاق القديم، نشر التطبيق على المتجرين، تسجيل فيديو تدريبي للوحة التحكم، تسليم الكود المصدري، بدء فترة الضمان."),
    ]
    for title, duration, body in phases:
        add_subheading(doc, f"المرحلة {title}")
        add_arabic_para_runs(doc, [
            {"text": "المدة: ", "bold": True, "color": GREEN},
            {"text": duration, "color": DARK},
        ], size=11, space_after=2)
        add_arabic_para_runs(doc, [
            {"text": "المخرجات: ", "bold": True, "color": GREEN},
            {"text": body, "color": DARK},
        ], size=11, space_after=8, line_spacing=1.5)

    add_callout(doc,
        "ملاحظة جوهرية",
        "يُعدّ اعتماد مخرجات أي مرحلة بالبريد الإلكتروني الرسمي أو الواتساب الرسمي "
        "اعتماداً كاملاً وملزماً، ولا يحق للطرف الأول التراجع عن الاعتماد بأثر رجعي.",
    )

    # ============== المادة 4: المدة الزمنية ==============
    add_heading_article(doc, "الرابعة", "المدة الزمنية")
    add_numbered_item(doc, 1, "مدة تنفيذ المشروع: اثنا عشر أسبوعاً متّصلاً (12 أسبوعاً) تبدأ من تاريخ سداد الدفعة الأولى واستلام جميع المواد والمحتوى من الطرف الأول، أيهما لاحقاً.")
    add_placeholder_para(doc, "تاريخ بدء التنفيذ المتوقع", "يوم/شهر/سنة")
    add_placeholder_para(doc, "تاريخ التسليم النهائي المتوقع", "يوم/شهر/سنة")
    add_numbered_item(doc, 2, "في حال تأخر الطرف الأول في تقديم المحتوى أو الاعتمادات أو الوصول إلى أي خدمة خارجية لازمة، يمتد الجدول الزمني تلقائياً بمدة مساوية لفترة التأخير، دون أي رسوم إضافية أو تبعات على الطرف الثاني.")
    add_numbered_item(doc, 3, "في حال طلب الطرف الأول أعمالاً خارج نطاق الملحق (أ) (تعديل نطاق)، يمتد الجدول الزمني بمدة إضافية يتفق عليها الطرفان كتابياً وفقاً للمادة الحادية عشرة.")
    add_numbered_item(doc, 4, "يقدّم الطرف الثاني تحديثاً أسبوعياً عن نسبة الإنجاز عبر القناة الرسمية للتواصل المتفق عليها.")

    # ============== المادة 5: القيمة المالية والدفعات ==============
    add_heading_article(doc, "الخامسة", "القيمة المالية والدفعات")
    add_arabic_para(doc,
        "اتّفق الطرفان على القيمة المالية للمشروع وجدول الدفعات على النحو التالي:",
        size=12, space_after=8, line_spacing=1.5)

    add_subheading(doc, "أولاً · القيمة الإجمالية")
    add_numbered_item(doc, 1, "القيمة الإجمالية للعقد: عشرون ألف ريال سعودي (20,000 ر.س) فقط لا غير.")
    add_numbered_item(doc, 2, "هذه القيمة نهائية وشاملة لجميع الضرائب والرسوم والأتعاب الإدارية والفنية، ولا تُضاف إليها أي ضريبة قيمة مضافة ولا أي مبالغ أخرى من أي نوع.")
    add_numbered_item(doc, 3, "العملة المعتمدة للسداد: الريال السعودي. ويجوز السداد بما يعادله بعملة أخرى وفق سعر الصرف يوم التحويل، على أن يتحمّل الطرف الأول كامل فروقات سعر الصرف ورسوم التحويل الدولي.")

    add_subheading(doc, "ثانياً · جدول الدفعات")
    # Payment table
    pay_table = doc.add_table(rows=7, cols=4)
    pay_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pay_table.autofit = False
    make_table_rtl(pay_table)
    set_table_borders(pay_table)
    widths = [Cm(2.5), Cm(7.5), Cm(3.5), Cm(2.0)]
    for i, w in enumerate(widths):
        pay_table.columns[i].width = w

    # Header
    headers = ["الدفعة", "المرحلة المرتبطة", "المبلغ (ر.س)", "النسبة"]
    for i, h in enumerate(headers):
        write_cell(pay_table.cell(0, i), h,
                   size=11, bold=True, color=DARK,
                   align=WD_ALIGN_PARAGRAPH.CENTER, shade=GREEN_BG)

    rows = [
        ("الأولى", "عند توقيع العقد وقبل بدء التصميم", "3,000", "15%"),
        ("الثانية", "عند اعتماد مخرجات التصميم", "5,000", "25%"),
        ("الثالثة", "عند اعتماد الواجهة الأمامية ولوحة التحكم", "5,000", "25%"),
        ("الرابعة", "عند اعتماد الخدمات الخلفية وربط الأنظمة", "5,000", "25%"),
        ("الخامسة", "عند اعتماد التسليم النهائي للموقع والتطبيق", "2,000", "10%"),
        ("الإجمالي", "—", "20,000", "100%"),
    ]
    for r_idx, (a, b, c, d) in enumerate(rows, start=1):
        is_total = (a == "الإجمالي")
        shade = "F8F8F8" if is_total else None
        write_cell(pay_table.cell(r_idx, 0), a,
                   size=11, bold=True,
                   color=GREEN if not is_total else DARK,
                   align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)
        write_cell(pay_table.cell(r_idx, 1), b, size=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)
        write_cell(pay_table.cell(r_idx, 2), c, size=11,
                   bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)
        write_cell(pay_table.cell(r_idx, 3), d, size=11,
                   bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_subheading(doc, "ثالثاً · بيانات الحساب البنكي للسداد")
    add_kv_line(doc, "اسم المستفيد", "AHMED ALI BASSIONI · أحمد علي بسيوني علي")
    add_kv_line(doc, "اسم البنك", "بنك الإسكندرية · Bank of Alexandria")
    add_kv_line(doc, "رقم الحساب", "311175442001")
    add_kv_line(doc, "رقم الآيبان (IBAN)", "EG230005301100000311175442001")
    add_kv_line(doc, "الدولة", "جمهورية مصر العربية")

    add_subheading(doc, "رابعاً · أحكام السداد")
    add_numbered_item(doc, 1, "تُسدّد كل دفعة في موعد أقصاه سبعة أيام عمل (7) من تاريخ إشعار الطرف الثاني الكتابي باكتمال المرحلة المرتبطة بها.")
    add_numbered_item(doc, 2, "يحق للطرف الثاني إيقاف العمل في المرحلة التالية حتى استلام الدفعة المستحقة عن المرحلة السابقة، ويمتد الجدول الزمني بمدة التأخر في السداد.")
    add_numbered_item(doc, 3, "تُعدّ الدفعات نهائية وغير قابلة للاسترداد بعد سدادها، باستثناء حالة الإنهاء الواردة في المادة الثانية عشرة.")
    add_numbered_item(doc, 4, "يلتزم الطرف الثاني بإصدار سند استلام مالي أو فاتورة لكل دفعة فور استلامها.")

    # ============== المادة 6: التزامات الطرف الأول ==============
    add_heading_article(doc, "السادسة", "التزامات الطرف الأول (العميل)")
    obligations_1 = [
        "سداد الدفعات في مواعيدها المنصوص عليها في المادة الخامسة دون تأخير.",
        "تزويد الطرف الثاني بجميع المحتوى النصي والصور والشعارات وأصول الهوية البصرية اللازمة لكل مرحلة، قبل بدء تنفيذها بمدة كافية.",
        "تقديم الاعتمادات الكتابية على مخرجات كل مرحلة خلال خمسة أيام عمل (5) من تاريخ تسليمها، وفي حال تأخر ذلك يمتد الجدول الزمني تلقائياً.",
        "فتح حساب Apple Developer بقيمة 99 دولاراً أمريكياً سنوياً، وحساب Google Play Console بقيمة 25 دولاراً أمريكياً لمرة واحدة، باسم الشركة وتزويد الطرف الثاني ببيانات الوصول قبل بدء المرحلة الرابعة.",
        "تحمّل جميع الاشتراكات والخدمات الخارجية المتكررة (مثل: خدمة الـ OTP من Unifonic، بوابة الدفع Tap أو Paymob ونسبها على كل عملية، استضافة Salam Cloud الشهرية، خدمة Sentry لرصد الأخطاء) بشكل مباشر ودون أي وسيط.",
        "تقديم بيانات الأعضاء النشطين الحاليين بصيغة قابلة للترحيل (CSV / SQL / JSON) قبل بدء المرحلة الثالثة.",
        "تعيين شخص اتصال موحّد (Single Point of Contact) للتنسيق مع الطرف الثاني طوال مدة المشروع، وإبلاغ الطرف الثاني كتابياً بأي تغيير فيه.",
        "الالتزام بصلاحيات وحقوق الطرف الثاني في الملكية الفكرية المنصوص عليها في المادة الثامنة.",
        "عدم تعديل الكود المصدري أو تكليف طرف ثالث بتعديله خلال فترة التطوير وفترة الضمان، إذ يُلغي ذلك التزامات الضمان الواردة في المادة العاشرة.",
    ]
    for i, item in enumerate(obligations_1, start=1):
        add_numbered_item(doc, i, item)

    # ============== المادة 7: التزامات الطرف الثاني ==============
    add_heading_article(doc, "السابعة", "التزامات الطرف الثاني (المطوّر)")
    obligations_2 = [
        "تنفيذ المشروع وفقاً للوثيقة التقنية المرفقة (الملحق أ) بكامل بنودها ومتطلباتها التفصيلية.",
        "تسليم مخرجات كل مرحلة في موعدها المحدد، مع تحمّل مسؤولية أي تأخير لا يعود لسبب يخصّ الطرف الأول أو قوة قاهرة.",
        "الالتزام بمعايير جودة الكود المعتمدة (Clean Code, Type Safety) ومعايير الأمان الأساسية وفق قائمة OWASP Top 10.",
        "تسليم الكود المصدري كاملاً عبر مستودع GitHub خاص يُنشأ باسم الطرف الأول وتنتقل ملكيته إليه بعد السداد النهائي.",
        "تسجيل فيديو تدريبي شامل للوحة التحكم الإدارية وتسليمه ضمن مخرجات المرحلة الخامسة.",
        "تقديم فترة ضمان فني مجانية لمدة ثلاثة أشهر (3) بعد الإطلاق الرسمي، وفق أحكام المادة العاشرة.",
        "الحفاظ على سرية بيانات المشروع والعميل وأعضائه وفق أحكام المادة التاسعة.",
        "تقديم تحديث أسبوعي للطرف الأول عن نسبة الإنجاز والعقبات إن وُجدت.",
        "عدم إسناد أي جزء من العمل لطرف ثالث دون الحصول على موافقة الطرف الأول الكتابية المسبقة، باستثناء التعاقدات المعتادة مع مزوّدي الخدمات السحابية والمكتبات مفتوحة المصدر.",
    ]
    for i, item in enumerate(obligations_2, start=1):
        add_numbered_item(doc, i, item)

    # ============== المادة 8: الملكية الفكرية ==============
    add_heading_article(doc, "الثامنة", "الملكية الفكرية")
    add_numbered_item(doc, 1, "تنتقل ملكية الكود المصدري الكامل والتصاميم والمخرجات النهائية ملكيةً تامةً للطرف الأول فور سداد كامل قيمة العقد البالغة عشرين ألف ريال سعودي (20,000 ر.س).")
    add_numbered_item(doc, 2, "قبل اكتمال السداد، يظل الكود المصدري والتصاميم في حيازة الطرف الثاني، ويُمنح الطرف الأول حق الاستخدام بقدر ما هو لازم للمراجعة والاعتماد.")
    add_numbered_item(doc, 3, "يحتفظ الطرف الثاني بحق ذكر المشروع ضمن معرض أعماله الشخصي (Portfolio) بشكل عام، دون الإفصاح عن أي بيانات سرية أو شيفرة برمجية تخصّ الطرف الأول.")
    add_numbered_item(doc, 4, "تراخيص المكتبات والأدوات مفتوحة المصدر المستخدمة في المشروع تبقى وفق شروط مالكها الأصلي، ويُعدّ الطرف الأول مالكاً لحق استخدامها داخل المشروع فقط.")
    add_numbered_item(doc, 5, "تراخيص الخطوط والأيقونات والصور التجارية -إن وُجدت- يتحمّل الطرف الأول تكلفتها وتجديدها بشكل مباشر.")

    # ============== المادة 9: السرية ==============
    add_heading_article(doc, "التاسعة", "السرية")
    add_numbered_item(doc, 1, "يلتزم كلا الطرفين بالحفاظ التام على سرية المعلومات التي يطّلع عليها بسبب تنفيذ هذا العقد، وعدم إفشائها أو استخدامها لأي غرض خارج نطاق العقد.")
    add_numbered_item(doc, 2, "تشمل المعلومات السرية -دون حصر- ما يلي: بيانات الأعضاء، النماذج التشغيلية والمالية للعميل، الكود المصدري، الخطط التشويقية والتسويقية، أي بيانات تجارية أو شخصية يطّلع عليها الطرف الثاني بحكم عمله.")
    add_numbered_item(doc, 3, "يستمر التزام السرية لمدة خمس سنوات ميلادية (5) من تاريخ انتهاء العقد أو إنهائه لأي سبب.")
    add_numbered_item(doc, 4, "لا تشمل السرية: المعلومات المتاحة للعموم قبل توقيع العقد، أو المعلومات التي يُطلب الإفصاح عنها بموجب أمر قضائي أو طلب نظامي رسمي.")
    add_numbered_item(doc, 5, "يلتزم الطرف الثاني بحماية بيانات الأعضاء وفقاً لنظام حماية البيانات الشخصية (PDPL) المعمول به في المملكة العربية السعودية في جميع مكونات النظام.")

    # ============== المادة 10: الضمان والمسؤولية ==============
    add_heading_article(doc, "العاشرة", "الضمان والمسؤولية والقوة القاهرة")
    add_subheading(doc, "أولاً · فترة الضمان")
    add_numbered_item(doc, 1, "يقدّم الطرف الثاني ضماناً فنياً مجانياً لمدة ثلاثة أشهر (3) من تاريخ الإطلاق الرسمي للمنصة والتطبيق.")
    add_numbered_item(doc, 2, "يشمل الضمان: إصلاح أي خطأ وظيفي أو أمني يظهر في نطاق العمل المعتمد، دون مقابل.")
    add_numbered_item(doc, 3, "لا يشمل الضمان: تطوير ميزات جديدة، تعديل نطاق العمل، أي خلل ناتج عن تدخّل خارجي على الكود من غير الطرف الثاني، تغيّر متطلبات منصات النشر (App Store / Google Play)، أو تحديثات على الخدمات الخارجية.")

    add_subheading(doc, "ثانياً · حدود المسؤولية")
    add_numbered_item(doc, 1, "لا يتحمّل الطرف الثاني بأي حال من الأحوال أي مسؤولية عن الأضرار غير المباشرة أو التبعية (فقدان الأرباح، فرص العمل، السمعة، البيانات).")
    add_numbered_item(doc, 2, "في حال ثبتت أي مسؤولية مباشرة على الطرف الثاني، فإن الحد الأقصى الإجمالي لها لا يتجاوز كامل قيمة هذا العقد (20,000 ر.س).")

    add_subheading(doc, "ثالثاً · القوة القاهرة")
    add_numbered_item(doc, 1, "يُعفى أي من الطرفين من المسؤولية عن التأخر في تنفيذ التزاماته إذا كان التأخر بسبب قوة قاهرة تخرج عن إرادته، مثل: الكوارث الطبيعية، الحروب، القرارات الحكومية، انقطاع البنية التحتية للإنترنت لمدة طويلة.")
    add_numbered_item(doc, 2, "يلتزم الطرف المتأثر بإخطار الطرف الآخر خلال سبعة أيام (7) من حدوث الواقعة، ويمتد الجدول الزمني بمدة مساوية لفترة القوة القاهرة.")
    add_numbered_item(doc, 3, "إذا استمرت القوة القاهرة لأكثر من ستين يوماً (60)، يحق لأي من الطرفين طلب فسخ العقد ودياً، مع تسوية الحقوق المالية وفق نسبة الإنجاز الفعلي.")

    # ============== المادة 11: تعديلات نطاق العمل ==============
    add_heading_article(doc, "الحادية عشرة", "تعديلات نطاق العمل (Change Orders)")
    add_numbered_item(doc, 1, "يُعتبر أي طلب من الطرف الأول لإضافة ميزة أو تعديل وظيفة خارج ما هو موصوف في الوثيقة التقنية (الملحق أ) نطاق عمل إضافي يخضع لأحكام هذه المادة.")
    add_numbered_item(doc, 2, "يقدّم الطرف الثاني للطرف الأول كتابياً تقديراً لتكلفة الإضافة والمدة الزمنية المطلوبة لتنفيذها خلال خمسة أيام عمل (5) من استلام الطلب.")
    add_numbered_item(doc, 3, "لا يُباشر الطرف الثاني تنفيذ أي تعديل أو إضافة إلا بعد موافقة الطرف الأول الكتابية الصريحة على التكلفة والمدة.")
    add_numbered_item(doc, 4, "تُسدّد تكلفة كل تعديل مُعتمَد بشكل منفصل عن قيمة العقد الأصلية، وفق جدول دفعات يُتفق عليه لكل تعديل.")
    add_numbered_item(doc, 5, "تُعدّ الميزات المدرجة في القسم الثاني عشر من الوثيقة التقنية «المراحل المستقبلية» خارج نطاق هذا العقد، وتُنفَّذ -إن طُلبت- باتفاق مستقل.")

    # ============== المادة 12: الإنهاء ==============
    add_heading_article(doc, "الثانية عشرة", "إنهاء العقد")
    add_subheading(doc, "أولاً · الإنهاء بالتراضي")
    add_numbered_item(doc, 1, "يحق للطرفين بالاتفاق المتبادل إنهاء هذا العقد في أي وقت، على أن يكون ذلك بمحرّر كتابي موقّع من الطرفين، مع تسوية الحقوق المالية وفق نسبة الإنجاز الفعلي.")

    add_subheading(doc, "ثانياً · الإنهاء بسبب الإخلال")
    add_numbered_item(doc, 1, "في حال إخلال أي من الطرفين بأي التزام جوهري من التزاماته في هذا العقد، يحق للطرف الآخر منحه إنذاراً كتابياً يحدد فيه الإخلال ويمنحه مهلة ثلاثين يوماً (30) لتصحيحه.")
    add_numbered_item(doc, 2, "إذا لم يصحّح الطرف المخل إخلاله خلال المهلة، يحق للطرف الآخر إنهاء العقد بمحرّر كتابي.")

    add_subheading(doc, "ثالثاً · أثر الإنهاء")
    add_numbered_item(doc, 1, "يلتزم الطرف الثاني بتسليم جميع المخرجات المنجزة فعلياً حتى تاريخ الإنهاء، وفق حالتها التقنية في ذلك التاريخ.")
    add_numbered_item(doc, 2, "يستحق الطرف الثاني كامل قيمة المراحل المنجزة والمعتمدة، إضافةً إلى نسبة 50% من قيمة المرحلة الجارية إن كان قد بدأ فيها.")
    add_numbered_item(doc, 3, "تنتقل ملكية المخرجات المسلَّمة فعلياً للطرف الأول بعد سداد المستحقات المذكورة في الفقرة السابقة.")
    add_numbered_item(doc, 4, "تظل التزامات السرية الواردة في المادة التاسعة سارية بعد الإنهاء لمدتها المنصوص عليها.")

    # ============== المادة 13: الإشعارات ==============
    add_heading_article(doc, "الثالثة عشرة", "الإشعارات والتواصل الرسمي")
    add_numbered_item(doc, 1, "تُعدّ الإشعارات والمراسلات والاعتمادات والإخطارات صحيحة ومنتجة لآثارها القانونية إذا أُرسلت عبر إحدى القناتين التاليتين:")
    add_bullet(doc, "البريد الإلكتروني الرسمي المعتمد لكل طرف.")
    add_bullet(doc, "تطبيق الواتساب على الرقم الرسمي المعتمد لكل طرف.")

    add_subheading(doc, "القنوات الرسمية المعتمدة")
    # Notice table
    nt = doc.add_table(rows=3, cols=3)
    nt.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_table_rtl(nt)
    set_table_borders(nt)
    for i, w in enumerate([Cm(4), Cm(6), Cm(5.5)]):
        nt.columns[i].width = w
    headers = ["الطرف", "البريد الإلكتروني", "الواتساب الرسمي"]
    for i, h in enumerate(headers):
        write_cell(nt.cell(0, i), h, size=11, bold=True, color=DARK,
                   align=WD_ALIGN_PARAGRAPH.CENTER, shade=GREEN_BG)
    # Row party 1 (placeholders)
    write_cell(nt.cell(1, 0), "الطرف الأول", size=11, bold=True, color=GREEN,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(nt.cell(1, 1), "[ يُحدد عند التوقيع ]", size=11, color=PLACEHOLDER,
               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(nt.cell(1, 2), "[ يُحدد عند التوقيع ]", size=11, color=PLACEHOLDER,
               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Row party 2
    write_cell(nt.cell(2, 0), "الطرف الثاني", size=11, bold=True, color=GREEN,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(nt.cell(2, 1), "hello@ahmedali.online", size=11,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(nt.cell(2, 2), "+201011648156", size=11,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_numbered_item(doc, 2, "يلتزم كل طرف بإبلاغ الطرف الآخر كتابياً بأي تغيير في قنوات التواصل الرسمية خلال ثلاثة أيام (3) من حدوث التغيير.")

    # ============== المادة 14: القانون الواجب التطبيق ==============
    add_heading_article(doc, "الرابعة عشرة", "القانون الواجب التطبيق")
    add_arabic_para(doc,
        "يخضع هذا العقد في تفسيره وتنفيذه لأحكام الأنظمة المعمول بها في "
        "المملكة العربية السعودية، بما في ذلك نظام التعاملات الإلكترونية، "
        "ونظام مكافحة الجرائم المعلوماتية، ونظام حماية البيانات الشخصية "
        "(PDPL)، وأي أنظمة أخرى ذات صلة.",
        size=12, space_after=6, line_spacing=1.55, indent_first=0.5)

    # ============== المادة 15: أحكام عامة ==============
    add_heading_article(doc, "الخامسة عشرة", "أحكام عامة")
    add_numbered_item(doc, 1, "يمثّل هذا العقد بملاحقه كامل الاتفاق بين الطرفين، ويلغي ويحلّ محلّ أي مفاوضات أو وعود أو مراسلات أو اتفاقيات سابقة بينهما تتعلق بموضوع العقد.")
    add_numbered_item(doc, 2, "لا يجوز تعديل أي بند من بنود هذا العقد إلا بمحرر كتابي مُلحَق به وموقّع من الطرفين، ويُعدّ هذا الملحق جزءاً منه.")
    add_numbered_item(doc, 3, "عدم استعمال أي طرف لأي حق من حقوقه أو تأخره في استعماله لا يُعدّ تنازلاً عنه ولا يفقده الحق في استعماله مستقبلاً.")
    add_numbered_item(doc, 4, "إذا تبيّن أن أحد بنود هذا العقد باطل أو غير قابل للتنفيذ لأي سبب، يبقى باقي بنود العقد سارياً ومنتجاً لآثاره الكاملة.")
    add_numbered_item(doc, 5, "اللغة المعتمدة للعقد وملاحقه: العربية الفصحى. وفي حال تُرجم العقد لأي لغة أخرى لأي غرض، تظلّ النسخة العربية هي المعتمدة عند أي تعارض.")
    add_numbered_item(doc, 6, "العناوين الواردة في صدر كل مادة وُضعت للتيسير فقط، ولا تؤثر في تفسير بنودها.")

    # ============== المادة 16: نسخ العقد والملاحق ==============
    add_heading_article(doc, "السادسة عشرة", "نسخ العقد وملاحقه")
    add_numbered_item(doc, 1, "حُرّر هذا العقد من نسختين أصليتين متطابقتين، بيد كل طرف نسخة واحدة للعمل بمقتضاها.")
    add_numbered_item(doc, 2, "يتضمن هذا العقد الملاحق التالية، وتُعدّ جزءاً لا يتجزأ منه:")
    add_bullet(doc, "الملحق (أ): الوثيقة التقنية الكاملة لنطاق العمل (12 قسماً، 22 متطلباً، 5 مراحل).")
    add_bullet(doc, "الملحق (ب): جدول الدفعات التفصيلي (مدمج في المادة الخامسة).")
    add_numbered_item(doc, 3, "أقرّ الطرفان بأنهما قرآ هذا العقد وملاحقه قراءةً تامة، وفهماه فهماً نافياً للجهالة، وقبلاه ووقّعاه بكامل أهليتهما المعتبرة شرعاً ونظاماً.")

    # ============== التوقيعات ==============
    doc.add_page_break()
    add_arabic_para(doc, "التوقيعات",
                    size=22, bold=True, color=GREEN,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=4)
    add_arabic_para(doc,
        "يُوقّع الطرفان على هذا العقد إقراراً منهما بقبول جميع ما ورد فيه:",
        size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24)

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    make_table_rtl(sig_table)
    sig_table.columns[0].width = Cm(8)
    sig_table.columns[1].width = Cm(8)

    def fill_sig_cell(cell, side_label, party_name, fields):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        shade_cell(cell, "FAFAFA")
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for s in ("top", "bottom", "left", "right"):
            b = OxmlElement(f"w:{s}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:color"), "30C280")
            tcBorders.append(b)
        tcPr.append(tcBorders)

        # Side label
        p0 = cell.paragraphs[0]
        set_rtl(p0)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(side_label)
        set_arabic_font(r0, size=10, bold=True, color=GREEN)

        # Party name
        p1 = cell.add_paragraph()
        set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(14)
        r1 = p1.add_run(party_name)
        set_arabic_font(r1, size=13, bold=True, color=DARK)

        # Fields
        for label, value, is_placeholder in fields:
            pf = cell.add_paragraph()
            set_rtl(pf)
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf.paragraph_format.space_after = Pt(6)
            pf.paragraph_format.line_spacing = 1.4
            lbl = pf.add_run(f"{label}: ")
            set_arabic_font(lbl, size=11, bold=True, color=DARK)
            val = pf.add_run(value)
            set_arabic_font(val, size=11,
                            color=PLACEHOLDER if is_placeholder else DARK,
                            bold=is_placeholder)

        # Signature line
        for _ in range(2):
            cell.add_paragraph().paragraph_format.space_after = Pt(0)
        psig = cell.add_paragraph()
        set_rtl(psig)
        psig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        psig.paragraph_format.space_before = Pt(20)
        psig.paragraph_format.space_after = Pt(2)
        rsig = psig.add_run("________________________________")
        set_arabic_font(rsig, size=11, color=DARK)
        psig2 = cell.add_paragraph()
        set_rtl(psig2)
        psig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        psig2.paragraph_format.space_after = Pt(8)
        rsig2 = psig2.add_run("التوقيع والختم")
        set_arabic_font(rsig2, size=10, bold=True, color=GRAY)

    fill_sig_cell(
        sig_table.cell(0, 0),
        "الطرف الأول · العميل",
        "شركة خطابة السعودية الأولى\nللتجارة (ذ.م.م)",
        [
            ("اسم الممثل المفوّض", "[ يُملأ ]", True),
            ("الصفة", "[ يُملأ ]", True),
            ("التاريخ", "[ يُملأ ]", True),
        ],
    )
    fill_sig_cell(
        sig_table.cell(0, 1),
        "الطرف الثاني · المطوّر",
        "أحمد علي بسيوني علي",
        [
            ("الصفة", "مطوّر برمجيات مستقل", False),
            ("الجنسية", "مصري", False),
            ("التاريخ", "[ يُملأ ]", True),
        ],
    )

    # ============== ملحق (أ) - مرجعية ==============
    doc.add_page_break()
    add_arabic_para(doc, "الملحق (أ)",
                    size=20, bold=True, color=GREEN,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_arabic_para(doc, "الوثيقة التقنية الكاملة لنطاق العمل",
                    size=14, bold=True, color=DARK,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_arabic_para(doc,
        "تُعدّ الوثيقة التقنية المرفقة بهذا العقد جزءاً لا يتجزأ منه، وتقرأ "
        "معه قراءةً واحدةً. وتشمل ما يلي:",
        size=12, space_after=10, line_spacing=1.55, indent_first=0.5)

    annex_a = [
        "القسم الأول: ملخص المشروع (الاسم، العميل، المطوّر، المدة، القيمة).",
        "القسم الثاني: الحزمة التقنية المعتمدة (Frontend, Backend, Mobile, DevOps, Security).",
        "القسم الثالث: خريطة صفحات المنصة (17 صفحة) ولوحة التحكم (9 شاشات) والتطبيق.",
        "القسم الرابع: رحلة المستخدم في 12 مرحلة، نظام الشات المراقب 4.2، نظام الدفع المعتمد 4.3.",
        "القسم الخامس: المراحل التفصيلية والمخرجات لكل مرحلة.",
        "القسم السادس: الترتيب الزمني (Gantt) للمشروع.",
        "القسم السابع: خطة ترحيل بيانات الأعضاء من النظام القديم.",
        "القسم الثامن: هيكل قاعدة البيانات وأهم الجداول والعلاقات.",
        "القسم التاسع: اثنان وعشرون متطلباً وميزة تفصيلية (المحفظة، الحسابات الراكدة، التحكم بحساب العضو، مراجعة التحديثات، سجل الزيارات، كشف التكرار، التنبيه بعدم التوافق، آلية تسجيل ونشر الحساب، البحث الآلي، حالات العضوية، مراقبة السلوك، إشعارات التواصل، كشف المتصلين، الملاحظات الداخلية، الدخول السريع، وسائل الدفع، صلاحيات الإدارة، دليل الاستخدام، التقدم للخطبة، تسجيل الزوار، نظام عروض الزواج، اتفاقية ما قبل الواتساب).",
        "القسم العاشر: الإقرارات والتعهدات المطلوبة من الأعضاء.",
        "القسم الحادي عشر: نماذج التسجيل التفصيلية للرجل والمرأة.",
        "القسم الثاني عشر: المراحل المستقبلية (خارج نطاق هذا العقد).",
    ]
    for i, item in enumerate(annex_a, start=1):
        add_numbered_item(doc, i, item)

    add_callout(doc,
        "نسخة الوثيقة التقنية",
        "النسخة المعتمدة المرجعية: الإصدار 1.1 · تاريخ الاعتماد: مايو 2026. "
        "أي تعديل لاحق على الوثيقة يجب أن يكون كتابياً وموقّعاً من الطرفين، "
        "ويُلحق بالعقد كملحق إضافي.",
    )

    # ============== END ==============
    output = "/Users/ahmedali/Code/portifolio/contracts/khattaba-contract-v1.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    path = build_document()
    print(f"OK: {path}")
    print(f"size: {os.path.getsize(path):,} bytes")
